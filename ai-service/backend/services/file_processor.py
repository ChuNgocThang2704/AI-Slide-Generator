"""Xử lý file: docx, pdf, txt"""
from pathlib import Path
from docx import Document
from typing import Union
import re
import asyncio
import logging

try:
    import fitz  # PyMuPDF
    _FITZ_AVAILABLE = True
except ImportError:
    _FITZ_AVAILABLE = False
    import pdfplumber  # fallback

logger = logging.getLogger(__name__)

class FileProcessor:
    """Xử lý các loại file input"""

    def _clean_extracted_text(self, text: str) -> str:
        """
        Clean common DOCX/PDF extraction artifacts và normalize text.
        - Gộp các dòng ngắn thành paragraph
        - Chỉ giữ line break khi thực sự cần thiết (paragraph break)
        """
        if not text:
            return ""
        
        # Normalize line breaks
        t = text.replace("\u00a0", " ")  # NBSP
        t = t.replace("_x000D_", " ")
        t = t.replace("\\r", " ")
        t = t.replace("\r\n", "\n").replace("\r", "\n")
        
        # Split into lines (giữ cả dòng rỗng để nhận biết ngắt đoạn thật)
        lines = [ln.strip() for ln in t.split("\n")]

        def is_heading_line(line: str) -> bool:
            if not line:
                return False
            words = line.split()
            return bool(
                re.match(r"^(CHƯƠNG|Chương|PHẦN|Phần|MỤC|Mục)\s+", line)
                or (
                    re.match(r"^(\d+(\.\d+)*[\.)]?\s+).+", line)
                    and len(line) <= 100
                    and len(words) <= 16
                )
                or re.match(r"^(Slide|Silde)\s+\d+\s*[:\-–]", line, re.IGNORECASE)
                or (line.isupper() and 4 <= len(line) <= 70)
                or (
                    line.endswith(":")
                    and len(line) <= 70
                    and len(words) <= 10
                    and line[:1].isupper()
                    and not re.search(r"[,\.!\?]", line[:-1])
                )
            )

        def is_bullet_line(line: str) -> bool:
            return bool(re.match(r"^(\-|\*|•|\u2022|\+|\d+[\.)])\s+", line))

        paragraphs = []
        current_para = []
        current_kind = "text"  # text | bullet | heading

        def flush_current():
            nonlocal current_para, current_kind
            if current_para:
                paragraphs.append(" ".join(current_para).strip())
            current_para = []
            current_kind = "text"

        for line in lines:
            if not line:
                flush_current()
                continue

            if is_heading_line(line):
                flush_current()
                paragraphs.append(line)
                current_kind = "heading"
                continue

            if is_bullet_line(line):
                flush_current()
                current_para = [line]
                current_kind = "bullet"
                continue

            # Dòng thường: gộp với paragraph hiện tại để tránh vỡ câu do PDF wrap
            if not current_para:
                current_para = [line]
                current_kind = "text"
            else:
                # Nếu đang trong bullet thì coi đây là phần tiếp nối bullet
                current_para.append(line)

        flush_current()

        # Merge các đoạn mở đầu bằng liên từ/chữ thường với đoạn trước để tránh cụt ý
        merged = []
        continuation_prefix = ("và ", "hoặc ", "nhưng ", "tuy ", "sau đó", "đồng thời", "từ đó")
        for para in paragraphs:
            p = (para or "").strip()
            if not p:
                continue
            p_lower = p.lower()
            prev = merged[-1] if merged else ""
            prev_ends_sentence = bool(re.search(r"[\.!\?:]$", prev.strip()))
            is_structured_line = bool(is_heading_line(p) or is_bullet_line(p))
            should_merge = bool(merged) and (
                p[:1].islower()
                or any(p_lower.startswith(pref) for pref in continuation_prefix)
                or (not prev_ends_sentence and not is_structured_line)
            )
            if should_merge:
                merged[-1] = f"{merged[-1]} {p}".strip()
            else:
                merged.append(p)

        paragraphs = merged
        
        # Join paragraphs với double newline (paragraph break)
        result = "\n\n".join(paragraphs)
        
        # Clean up spaces
        result = re.sub(r" +", " ", result)  # Multiple spaces -> single space
        result = re.sub(r"\n{3,}", "\n\n", result)  # Multiple newlines -> double
        
        return result.strip()
    
    async def process_file(self, file_path: Union[str, Path]) -> str:
        """Đọc và trích xuất text từ file"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension == ".docx":
            return await self._process_docx(file_path)
        elif extension == ".pdf":
            return await self._process_pdf(file_path)
        elif extension == ".txt":
            return await self._process_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Đọc nội dung DOCX (paragraphs + tables) theo luồng sync."""
        doc = Document(file_path)
        paragraphs = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            # Clean artifacts
            text = text.replace("\u00a0", " ")
            text = text.replace("_x000D_", " ")
            text = re.sub(r" +", " ", text)

            # Phát hiện style (Heading 1, 2, 3...)
            style_name = paragraph.style.name if paragraph.style else ""

            # Thêm markup để LLM hiểu cấu trúc
            if "Heading 1" in style_name or "Title" in style_name:
                text = f"# {text}"  # Heading cấp 1
            elif "Heading 2" in style_name:
                text = f"## {text}"  # Heading cấp 2
            elif "Heading 3" in style_name:
                text = f"### {text}"  # Heading cấp 3
            elif "Heading" in style_name:
                text = f"#### {text}"  # Heading cấp khác

            paragraphs.append(text)

        # Đọc thêm text trong bảng, giữ cấu trúc dạng Markdown để LLM hiểu cột/hàng.
        for table_idx, table in enumerate(doc.tables, start=1):
            rows = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    val = re.sub(r"\s+", " ", cell.text.replace("\u00a0", " ")).strip()
                    cells.append(val)
                if any(cells):
                    rows.append(cells)
            if not rows:
                continue

            max_cols = max(len(r) for r in rows)
            normalized = [r + [""] * (max_cols - len(r)) for r in rows]
            table_lines = [f"Table {table_idx}:"]
            header = normalized[0]
            table_lines.append("| " + " | ".join(header) + " |")
            table_lines.append("| " + " | ".join("---" for _ in header) + " |")
            for row in normalized[1:]:
                table_lines.append("| " + " | ".join(row) + " |")
            paragraphs.append("\n".join(table_lines))

        # Join với double newline để phân biệt paragraph breaks
        return "\n\n".join(paragraphs)

    async def _process_docx(self, file_path: Path) -> str:
        """Xử lý file DOCX - đọc styles (Heading 1, 2, 3...) để phân tích cấu trúc"""
        result = await asyncio.to_thread(self._extract_docx_text, file_path)
        return self._clean_extracted_text(result)

    def _validate_extraction_quality(self, text: str) -> bool:
        """Phát hiện text bị nối liền do PDF extract lỗi (ví dụ: 'MỤCLỤC2LỜICẢMƠN3').

        Trả về False nếu tỷ lệ 'từ dài bất thường' vượt ngưỡng 15%.
        """
        words = text.split()
        if not words:
            return False
        # Từ > 25 ký tự và không phải URL → bất thường
        long_word_count = sum(
            1 for w in words
            if len(w) > 25 and not w.startswith(("http", "www", "ftp"))
        )
        ratio = long_word_count / max(1, len(words))
        if ratio > 0.15:
            logger.warning(
                "[pdf_extract] %.0f%% abnormal long words detected — "
                "PDF may have lost whitespace (e.g. pdfplumber layout bug). "
                "Consider a text-layer PDF or DOCX.",
                ratio * 100,
            )
            return False
        return True

    def _extract_pdf_text_fitz(self, file_path: Path) -> str:
        """Dùng PyMuPDF (fitz) — giữ đúng khoảng trắng và cấu trúc dòng."""
        doc = fitz.open(str(file_path))
        pages: list[str] = []
        for page_number, page in enumerate(doc, start=1):
            # "text" mode + preserve_whitespace giữ đúng layout
            text = page.get_text("text")  # type: ignore[arg-type]
            if text and text.strip():
                pages.append(f"[[SOURCE_PAGE:{page_number}]]\n{text}")
        doc.close()
        return "\n\n".join(pages)

    def _extract_pdf_text_plumber(self, file_path: Path) -> str:
        """Fallback: dùng pdfplumber khi fitz không available."""
        import pdfplumber  # lazy import
        text_content: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                text = page.extract_text(x_tolerance=2, y_tolerance=3)
                if text:
                    text_content.append(f"[[SOURCE_PAGE:{page_number}]]\n{text}")
        return "\n\n".join(text_content)

    def _extract_pdf_text(self, file_path: Path) -> str:
        """Đọc text PDF theo luồng sync — ưu tiên fitz, fallback pdfplumber."""
        if _FITZ_AVAILABLE:
            return self._extract_pdf_text_fitz(file_path)
        return self._extract_pdf_text_plumber(file_path)

    async def _process_pdf(self, file_path: Path) -> str:
        """Xử lý file PDF với validate chất lượng extract."""
        raw_text = await asyncio.to_thread(self._extract_pdf_text, file_path)
        cleaned = self._clean_extracted_text(raw_text)
        if len(cleaned.strip()) < 100:
            raise ValueError(
                "PDF này có vẻ là file scan hoặc không có text layer nên không trích xuất được nội dung. "
                "Vui lòng dùng PDF có text layer, DOCX hoặc TXT."
            )
        # Validate chất lượng — cảnh báo nhưng không chặn (vẫn trả về để pipeline cố gắng xử lý)
        self._validate_extraction_quality(cleaned)
        return cleaned

    async def _process_txt(self, file_path: Path) -> str:
        """Xử lý file TXT"""
        def _read_text() -> str:
            for enc in ("utf-8", "utf-8-sig"):
                try:
                    with open(file_path, "r", encoding=enc) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()

        raw = await asyncio.to_thread(_read_text)
        return self._clean_extracted_text(raw)
